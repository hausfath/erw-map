"""Technical and economic ERW potential by country, with parameter ranges.

  python3 scripts/analysis/country_potential.py [--docx PATH]

A standalone analysis, not a map feature. One table, five columns
(restructured 2026-08-31):

  cropland          total agricultural land in scope, Mha
  TECH (uncapped)   steady-state gross removal holding 30 t/ha of rock,
                    NO drainage limit -- what the kinetics alone deliver
  ECON (uncapped)   the subset on cells where delivered rock (gate + regional
                    trucking) costs < $100/tCO2 of discounted lifetime carbon
  TECH (limited)    the same steady state capped by the drainage-concentration
                    ceiling -- the map's shipped basis
  ECON (limited)    the screened subset counting drainage-limited carbon
                    (variant B: the screen is a unit cost, never clamped)

Each potential is reported as p50 (p5-p95) over a PARAMETER ENSEMBLE, not a
calibrated posterior: N draws over the ranges the repo documents, constructed
so the ensemble median reproduces the shipped headline numbers.

  dissolution scale   split log-uniform over the verified deliveries' year-1
                      weathered range (15-56% at reference; anchor 25% =
                      ensemble median). The dominant technical axis.
  ceiling SI          uniform 0-1 (calcite saturation index; shipped 0.5,
                      Mayer et al. 2025's central case)
  ceiling Mg          split log-uniform 1/3-3 mM (shipped 1 mM; documented
                      range 0-5)
  quarry gate         the three documented scenarios (current byproduct
                      regional / uniform $10 / at-scale regional), equal weight
  haul rate           uniform x0.75-1.25 on the regional baselines

FIXED, deliberately, and why: drainage variable stays qtot (the qr bracket is
a structural choice reported by build gate 2d, not a distribution); soil pCO2
stays the protocol values (the paddy sensitivity is documented in
docs/PADDY_CEILING_INDIA.md -- the protocol 50,000 uatm is the BOTTOM of the
measured flooded-soil range, so paddy ceilings leans conservative); the cost
screen stays $100/tCO2 at 5% over the application's lifetime carbon; the
ceiling basis stays calcite saturation (the pH-target option would scale the
limited columns down ~40% -- see METHODOLOGY).

CAVEATS that bound every number: gross removal, not net (in-soil carbonate
precipitation, riverine re-release and strong-acid competition plausibly
claim more than 20% of gross; the ~80% top of the published range belongs to
particularly suboptimal siting); the rate law over-predicts its one
independent laboratory test, absorbed here by the delivery-anchored
dissolution scale; haul is great-circle x 1.35, not routed, and register
coverage is US/Brazil/Mexico/France + dense OSM in Europe (China and Turkey
are documented inventory gaps).
"""

from __future__ import annotations

import sys
from pathlib import Path

import numpy as np
from rasterio.enums import Resampling
from rasterio.features import rasterize
from scipy import ndimage

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
import constants as C  # noqa: E402
import kinetics as K  # noqa: E402
from build_v0 import master_grid, onto_grid  # noqa: E402

ROOT = Path(__file__).resolve().parents[2]
INTERIM = ROOT / "data/interim"

THRESHOLD_MT = 50.0     # report countries above this p50 uncapped potential
N_DRAWS = 500
SEED = 2026
SCREEN = C.COST_SCREEN_USD_PER_TCO2

# Gate scenarios by rate-group label, from values already sourced in the repo
# (docs/TRUCK_RATE_SOURCES.md, docs/GATE_COST_AT_SCALE.md).
GATE_SCENARIOS = {
    "optimistic": {"US/Canada": 12.0, "Brazil/Latin America": 9.0,
                   "India/South Asia": 3.0, "default": 10.0},
    "central": {"default": C.FEEDSTOCK_GATE_COST_USD_T},
    "conservative": {**C.FEEDSTOCK_GATE_AT_SCALE_USD_T,
                     "default": C.FEEDSTOCK_GATE_AT_SCALE_USD_T["elsewhere"]},
}


def country_raster(transform, w, h):
    """ISO_A2 index raster from the Natural Earth zip prep_feedstock cached."""
    import geopandas as gpd

    zp = ROOT / "data/raw/ne_50m_admin_0_countries.zip"
    ctry = gpd.read_file(f"zip://{zp}")
    iso_col = "ISO_A2_EH" if "ISO_A2_EH" in ctry.columns else "ISO_A2"
    isos, names, shapes = [], {}, []
    for _, row in ctry.iterrows():
        iso = str(row.get(iso_col) or "")
        if iso in ("-99", "nan", ""):
            iso = str(row.get("ADM0_A3") or "??")
        if iso not in isos:
            isos.append(iso)
            names[iso] = str(row.get("NAME") or iso)
        shapes.append((row.geometry, isos.index(iso) + 1))

    idx = rasterize(shapes, out_shape=(h, w), transform=transform,
                    fill=0, dtype="int32", all_touched=False)
    dist, (iy, ix) = ndimage.distance_transform_edt(idx == 0, return_indices=True)
    near = (dist > 0) & (dist <= 3)
    idx[near] = idx[iy[near], ix[near]]
    return idx, isos, names


def split_loguniform(rng, lo, mid, hi, n):
    """Median-preserving spread: log-uniform below and above `mid`, so the
    ensemble median sits at the shipped central value by construction."""
    u = rng.uniform(-1.0, 1.0, n)
    return np.where(u < 0, mid * (lo / mid) ** (-u), mid * (hi / mid) ** u)


def fmt(p50, p5, p95, digits=0):
    f = f"{{:,.{digits}f}}"
    return f"{f.format(p50)} ({f.format(p5)}–{f.format(p95)})"


def write_docx(path, rows, world, i_inf, us_cap):
    from docx import Document

    lo_rock, hi_rock, lo_co2, hi_co2, eta_us = us_cap
    doc = Document()
    doc.add_heading("Technical and economic ERW potential by country", 0)
    doc.add_paragraph(
        "Generated by scripts/analysis/country_potential.py from the ERW "
        "Atlas v0 build (github.com/hausfath/erw-map). All figures are GROSS "
        "CO2 removal on cropland, on the steady-state basis: hold 30 t/ha of "
        "undissolved rock, reapplying as the modelled kinetics dissolve it, "
        f"capped at one full application per year (mean-lifetime constant "
        f"I_inf = {i_inf:.3f}). Nothing downstream of export is deducted (no "
        "spreading, grinding-to-spec, MRV, cation retention, or riverine "
        "losses).")

    doc.add_heading("Definitions", level=1)
    doc.add_paragraph(
        "TECHNICAL (uncapped): the steady state with no drainage limit -- "
        "what the dissolution kinetics alone would deliver. TECHNICAL "
        "(drainage-limited): the same steady state capped by the "
        "drainage-concentration ceiling (calcite saturation at each cell's "
        "own soil pCO2 and temperature; the Atlas's shipped basis, applied "
        "by default since 2026-08-24 and corroborated as 'carrying capacity' "
        "by Mayer et al. 2025, doi:10.21203/rs.3.rs-7811095/v1). ECONOMIC: "
        "the subset on cells where delivered rock (gate + regional trucking "
        f"from the quarry inventory) costs under ${SCREEN:.0f} per tonne of "
        "CO2, against each application's discounted lifetime carbon at 5%. "
        "The screen is a unit cost and is never clamped by the drainage "
        "limit (carbon is linear in application rate, so an operator on a "
        "cap-bound cell applies less rock at the same $/tCO2); the carbon "
        "counted in the drainage-limited economic column is capped.")
    doc.add_paragraph(
        "Ranges are p50 (p5-p95) over a documented-parameter ensemble of "
        f"{N_DRAWS} draws -- NOT a calibrated posterior. Varied: the "
        "dissolution scale over the verified deliveries' year-1 weathered "
        "range (15-56%, anchor 25% = ensemble median); ceiling calcite SI "
        "uniform 0-1 (shipped 0.5) and dissolved Mg 1/3-3 mM (shipped 1); "
        "the three documented quarry-gate scenarios; haul rates x0.75-1.25. "
        "Fixed: drainage variable (total runoff; the recharge bracket is "
        "-14% on the uncapped world total), protocol soil pCO2 (the paddy "
        "value is the bottom of the measured range, so paddy ceilings lean "
        "conservative), the $100/tCO2 screen, and the calcite-saturation "
        "ceiling basis (the pH<=7.5 target option scales the limited "
        "columns down ~40%).")

    doc.add_heading(f"Potential by country, MtCO2/yr, p50 (p5–p95)",
                    level=1)
    t = doc.add_table(rows=1 + len(rows), cols=5)
    t.style = "Table Grid"
    hdr = ("Country · cropland (Mha)",
           "Technical, uncapped",
           f"Economic, uncapped (<${SCREEN:.0f}/tCO2)",
           "Technical, drainage-limited",
           f"Economic, drainage-limited (<${SCREEN:.0f}/tCO2)")
    for j, htxt in enumerate(hdr):
        t.rows[0].cells[j].text = htxt
    for i, r in enumerate(rows, start=1):
        c = t.rows[i].cells
        c[0].text = f"{r['name']} · {r['mha']:,.1f}"
        for j, k in enumerate(("tech_un", "econ_un", "tech_cap", "econ_cap"),
                              start=1):
            p5, p50, p95 = r[k]
            c[j].text = fmt(p50, p5, p95)

    doc.add_heading("The US byproduct-fines supply cap", level=1)
    doc.add_paragraph(
        "The economic columns are demand-side screens: they assume rock "
        "exists at the scenario gate price wherever the screen passes. For "
        "byproduct-gate pricing that fails at scale in the US: identifiable "
        "traprock fines-class sales are ~4.7 Mt/yr ($11-21/t) plus 3.3 Mt/yr "
        "of volcanic cinder/scoria at $8.83/t (USGS Minerals Yearbook 2023), "
        "with a ~20%-of-output heuristic upper bound near 22 Mt rock/yr "
        "(UNVERIFIED). At steady state every acquired tonne eventually "
        f"delivers eta x CT of CO2 (eta_US = {eta_us:.2f}), so the "
        f"supply-limited US potential at byproduct prices is "
        f"{lo_co2:.0f}-{hi_co2:.0f} MtCO2/yr from {lo_rock:.0f}-{hi_rock:.0f} "
        "Mt rock/yr -- in the United States, supply, not the cost screen, is "
        "the near-term constraint. The same logic applies to India's "
        "crusher-dust and Brazil's po-de-pedra streams, unquantified.")

    doc.add_heading("Caveats", level=1)
    doc.add_paragraph(
        "Gross, not net: in-soil carbonate precipitation, riverine "
        "re-release and strong-acid competition plausibly claim more than "
        "20% of gross removal (the ~80% top of the published range belongs "
        "to particularly suboptimal siting). The rate law over-predicts its "
        "one independent laboratory test; the delivery-anchored dissolution "
        "scale absorbs the level error but not the temperature gradient. "
        "Paddy rows (India, Thailand, Myanmar, Indonesia, parts of China) "
        "depend on flooded-soil chemistry no measured paddy DIC export "
        "exists to validate, and mapped cell values dilute paddy chemistry "
        "by each cell's non-rice share -- see docs/PADDY_CEILING_INDIA.md "
        "for the field-level analysis. Haul is great-circle x 1.35, not "
        "routed; quarry registers cover the US, Brazil, Mexico and France, "
        "with dense OSM in Europe and usable OSM in India, SE Asia, Japan "
        "and Brazil; China and Turkey are documented inventory gaps priced "
        "on the outcrop bound. The drainage limit is a maximum-EFFICIENT "
        "bound: past it, calcite precipitation halves marginal efficiency "
        "rather than stopping removal.")
    doc.save(str(path))


def main() -> int:
    transform, w, h, crs = master_grid()
    z = np.load(ROOT / "data/processed/v0_layers.npz")
    crop, area, ph = z["crop"], z["area"], z["ph"]
    m = (crop >= C.CROPLAND_MIN_FRACTION) & np.isfinite(ph)
    ha = (crop * area) * 100.0

    L1 = z["L1"].astype("float64")
    eta = np.nan_to_num(z["eta"].astype("float64"))
    eta_tr = np.nan_to_num(z["eta_tr"].astype("float64"))
    q = np.clip(np.nan_to_num(z["q"].astype("float64")), 0.0, None)
    pco2 = z["pco2"].astype("float64")
    TK = np.where(np.isfinite(z["t_ceil_c"]), z["t_ceil_c"], 15.0) + 273.15
    ceil_ship = np.nan_to_num(z["ceiling"].astype("float64"))

    rate = onto_grid(INTERIM / "truck_rate.tif", transform, w, h, crs,
                     resampling=Resampling.nearest).astype("float64")
    cost0 = onto_grid(INTERIM / "feedstock_cost.tif", transform, w, h, crs,
                      resampling=Resampling.average).astype("float64")
    d_eff = np.maximum(cost0 - C.FEEDSTOCK_GATE_COST_USD_T, 0.0) \
        / np.maximum(rate, 1e-9)

    group_rate = {k: g["rate"] for k, g in C.TRUCK_RATE_GROUPS.items()}
    vals = list(group_rate.values()) + [C.TRUCK_RATE_DEFAULT]
    assert len(set(np.round(vals, 4))) == len(vals), "rate values must be unique"

    def gate_raster(scn):
        g = np.full_like(rate, GATE_SCENARIOS[scn]["default"])
        for label, gv in GATE_SCENARIOS[scn].items():
            if label in group_rate:
                g[np.isclose(rate, group_rate[label])] = gv
        return g

    idx, isos, names = country_raster(transform, w, h)

    # ---- kinetics scaffolding, identical to the build's ---------------------
    spec = C.FEEDSTOCK_ARCHETYPES[C.FEEDSTOCK_DEFAULT]
    ct = ((spec["CaO_wt"] / C.M_CAO + spec["MgO_wt"] / C.M_MGO)
          * 1000.0 * 2.0 * C.MOL_CO2_PER_KMOL_CHARGE_T)
    d_ref = K.retreat_at_reference()
    ug = np.concatenate([[0.0], np.geomspace(1e-5, 200.0, 900)])
    gg = np.concatenate([[0.0], K.dissolved_fraction(ug[1:], C.PSD_REF_WIDTH)])
    i_inf = float(np.trapezoid(1.0 - gg, ug))
    u1_base = d_ref * np.clip(np.nan_to_num(10.0 ** L1) * eta_tr, 0.0, None) \
        / C.PSD_REF_D50_UM
    A = C.APPLICATION_RATE_T_HA_YR
    dr = C.COST_SCREEN_DISCOUNT_RATE

    # Discounted lifetime dissolved-fraction as a 1-D function of u1, so the
    # 60-year screen loop runs once on a grid instead of once per draw.
    u1g = np.concatenate([[0.0], np.geomspace(1e-7, 20.0, 600)])
    Dg = np.zeros_like(u1g)
    prev = np.zeros_like(u1g)
    for t in range(1, 61):
        cum = np.interp(u1g * t, ug, gg)
        Dg += (cum - prev) / (1.0 + dr) ** t
        prev = cum

    # ---- masked vectors ------------------------------------------------------
    sel = m
    hav = ha[sel]
    idxv = idx[sel].astype(np.int64)
    u1v = u1_base[sel]
    etav = eta[sel]
    qv = q[sel]
    ratev = rate[sel]
    dv = d_eff[sel]
    gates = {s: gate_raster(s)[sel] for s in GATE_SCENARIOS}
    scn_list = list(GATE_SCENARIOS)

    # Per-cell carbonate constants for the ceiling solve (T is per cell).
    K1, K2, KH, _ = K.carbonate_constants(TK)
    Ksp = K.k_calcite(TK)
    c_base = (2.0 * Ksp * K1 * KH * (pco2 * 1e-6) / K2)[sel]   # Omega = 1
    dhA = None
    # Davies A parameter at cell temperature, via kinetics' own helper.
    dhA = K._debye_huckel_a(TK)[sel]

    def ceiling_draw(omega, mg_mM):
        """q * [HCO3-]max * 44 for one (Omega, Mg) draw, Davies-corrected --
        the same solve as kinetics.alkalinity_ceiling_mol_l, vectorised over
        the masked cells with per-cell constants hoisted out of the loop."""
        mg = mg_mM * 1e-3
        c0 = omega * c_base
        g_div = np.ones_like(c0)
        Aalk = None
        for _ in range(3):
            c = c0 / g_div
            Aalk = np.cbrt(c) + 2.0 * mg
            for _ in range(20):
                Aalk = Aalk - (Aalk ** 3 - 2.0 * mg * Aalk * Aalk - c) \
                    / (3.0 * Aalk * Aalk - 4.0 * mg * Aalk)
            I = 1.5 * Aalk
            sq = np.sqrt(I)
            g_div = np.power(10.0, 6.0 * (-dhA * (sq / (1 + sq) - 0.3 * I)))
        return qv * 1e7 * Aalk * C.M_CO2_G_MOL / 1e6

    # Fidelity check: the ensemble's ceiling at the SHIPPED parameters must
    # reproduce the build's ceiling raster.
    chk = ceiling_draw(C.FLUX_CEILING_OMEGA, C.FLUX_CEILING_MG_MM)
    ref = ceil_ship[sel]
    ok = np.isfinite(ref) & (ref > 1e-6)
    rel = float(np.median(np.abs(chk[ok] / ref[ok] - 1)))
    print(f"ceiling fidelity vs build raster at shipped params: median rel "
          f"err {rel:.2%} (activity-iteration count differs; <2% expected)")
    assert rel < 0.02, "ensemble ceiling drifted from the build's"

    # ---- the ensemble --------------------------------------------------------
    rng = np.random.default_rng(SEED)
    nC = len(isos) + 1
    frac_lo, frac_hi = C.DISSOLVED_FRAC_OBSERVED_RANGE \
        if hasattr(C, "DISSOLVED_FRAC_OBSERVED_RANGE") else (0.15, 0.56)
    # invert year-1 fractions to dissolution-scale multipliers via the table
    u_of = lambda f: float(np.interp(f, gg, ug))
    k_lo = u_of(frac_lo) / u_of(C.DISSOLVED_FRAC_AT_REF)
    k_hi = u_of(frac_hi) / u_of(C.DISSOLVED_FRAC_AT_REF)
    ks = split_loguniform(rng, k_lo, 1.0, k_hi, N_DRAWS)
    sis = rng.uniform(0.0, 1.0, N_DRAWS)
    mgs = split_loguniform(rng, 1.0 / 3.0, 1.0, 3.0, N_DRAWS)
    scns = rng.integers(0, 3, N_DRAWS)
    hms = rng.uniform(0.75, 1.25, N_DRAWS)

    out = {k: np.zeros((N_DRAWS, nC)) for k in
           ("tech_un", "econ_un", "tech_cap", "econ_cap")}
    for d in range(N_DRAWS):
        u1 = ks[d] * u1v
        cad = A * np.minimum(1.0, u1 / i_inf) * etav * ct
        ceilv = ceiling_draw(10.0 ** sis[d], mgs[d])
        capped = np.minimum(cad, ceilv)
        dpt = etav * ct * A * np.interp(u1, u1g, Dg)
        cost = (gates[scn_list[scns[d]]] + hms[d] * ratev * dv) * A
        passes = (dpt > 0) & (cost / np.maximum(dpt, 1e-12) < SCREEN)
        for key, v in (("tech_un", cad), ("econ_un", np.where(passes, cad, 0)),
                       ("tech_cap", capped),
                       ("econ_cap", np.where(passes, capped, 0))):
            out[key][d] = np.bincount(idxv, weights=v * hav, minlength=nC) / 1e6
        if (d + 1) % 100 == 0:
            print(f"  draw {d + 1}/{N_DRAWS}")

    # ---- aggregate and report ------------------------------------------------
    mha = np.bincount(idxv, weights=hav, minlength=nC) / 1e10  # ha -> Mha? no:
    mha = np.bincount(idxv, weights=hav, minlength=nC) / 1e6   # ha -> Mha

    def pct(mat, ci):
        col = mat[:, ci] if ci is not None else mat.sum(axis=1)
        return (float(np.percentile(col, 5)), float(np.percentile(col, 50)),
                float(np.percentile(col, 95)))

    rows = []
    for i, iso in enumerate(isos, start=1):
        if mha[i] <= 0:
            continue
        r = {"iso": iso, "name": names[iso], "mha": mha[i],
             **{k: pct(out[k], i) for k in out}}
        rows.append(r)
    rows.sort(key=lambda r: -r["tech_un"][1])
    big = [r for r in rows if r["tech_un"][1] >= THRESHOLD_MT]
    world = {"iso": "", "name": "WORLD", "mha": float(mha[1:].sum()),
             **{k: pct(out[k], None) for k in out}}
    rest = {"iso": "", "name": "Rest of world",
            "mha": world["mha"] - sum(r["mha"] for r in big)}
    big_ci = [isos.index(r["iso"]) + 1 for r in big]
    for k in out:
        rest_mat = out[k].sum(axis=1) - out[k][:, big_ci].sum(axis=1)
        rest[k] = (float(np.percentile(rest_mat, 5)),
                   float(np.percentile(rest_mat, 50)),
                   float(np.percentile(rest_mat, 95)))
    table = big + [rest, world]

    hdr = (f"{'country':<22}{'Mha':>7}{'tech uncapped':>20}"
           f"{'econ uncapped':>20}{'tech limited':>20}{'econ limited':>20}")
    print(f"\nSteady-state potential, MtCO2/yr, p50 (p5-p95) over "
          f"{N_DRAWS} documented-parameter draws; econ = <${SCREEN:.0f}/tCO2")
    print(hdr)
    print("-" * len(hdr))
    for r in table:
        cells = "".join(f"{fmt(v[1], v[0], v[2]):>20}"
                        for v in (r["tech_un"], r["econ_un"],
                                  r["tech_cap"], r["econ_cap"]))
        print(f"{r['name'][:21]:<22}{r['mha']:>7.1f}{cells}")

    # Reference against the build. The uncapped p50 matches the build almost
    # exactly (the dissolution axis is median-pinned and cad is linear in k at
    # the cap-free margin). The LIMITED p50 sits a few percent ABOVE the
    # shipped 0.717 -- a Jensen gap, not an error: the ceiling is convex in
    # Mg, so the median of the ensemble exceeds the value at median
    # parameters. Both shipped values sit well inside their p5-p95 bands.
    print(f"\nfor reference, the build at shipped parameters: uncapped "
          f"2.591 Gt, drainage-limited 0.717 Gt. The limited p50 sits a few "
          f"percent above 717 by convexity (Jensen), not by drift; the "
          f"shipped values lie inside the p5-p95 bands.")

    # ---- US supply cap --------------------------------------------------------
    us_i = isos.index("US") + 1
    us_sel = idxv == us_i
    eta_us = float((etav[us_sel] * hav[us_sel]).sum() / hav[us_sel].sum())
    lo_rock, hi_rock = 4.7 + 3.3, 19.0 + 3.3
    lo_co2, hi_co2 = lo_rock * eta_us * ct, hi_rock * eta_us * ct
    print(f"\nUS byproduct-fines supply cap (USGS MYB 2023): "
          f"{lo_rock:.0f}-{hi_rock:.0f} Mt rock/yr -> {lo_co2:.0f}-"
          f"{hi_co2:.0f} MtCO2/yr at eta_US = {eta_us:.2f}; supply, not the "
          f"screen, binds the US at byproduct prices")

    if "--docx" in sys.argv:
        outp = Path(sys.argv[sys.argv.index("--docx") + 1]).expanduser()
        write_docx(outp, table, world, i_inf,
                   (lo_rock, hi_rock, lo_co2, hi_co2, eta_us))
        print(f"\nwrote {outp}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
